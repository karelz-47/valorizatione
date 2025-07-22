# streamlit_insurance_letter.py

# ---- Imports ------------------------------------------------------------
import locale
from datetime import date, timedelta
from io import BytesIO
import re                          
import pandas as pd
import streamlit as st
from babel.numbers import format_currency
from docx import Document            # ← ADD THIS LINE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt   #  ← add Pt
from typing import List

# -------------------------------------------------------------------------
#  LOCALE SETUP
# -------------------------------------------------------------------------
try:
    locale.setlocale(locale.LC_ALL, "it_IT.utf8")
except locale.Error:
    pass  # Streamlit Cloud may lack the locale, but Babel still formats OK

# ╔══════════════════════════════════════════════════════════════════════╗
# ║  HARD‑CODED CONFIGURATION                                            ║
# ╚══════════════════════════════════════════════════════════════════════╝

ITEM_CONFIG = {
    # ────────────────  Table T1  ────────────────
    "Acquisition cost deduction from regular premium": {
        "label": "Costi di emissione e gestione",
        "table": "T1",
        "pos": 1,
    },
    "Contract fee deduction from regular premium": {
        "label": "Costi di emissione e gestione",
        "table": "T1",
        "pos": 1,                 # same position; rows will merge
    },
    "Acquisition cost deduction from single premium": {
        "label": "Costi di emissione e gestione",
        "table": "T1",
        "pos": 1,
    },
    "Contract fee deduction from single premium": {
        "label": "Costi di emissione e gestione",
        "table": "T1",
        "pos": 1,
    },

    "Administrative deduction": {
        "label": "Costi di caricamento",
        "table": "T1",
        "pos": 2,
    },

    "Investment deduction": {
        "label": "Costi di investimento",
        "table": "T1",
        "pos": 3,
    },
    "Investment deduction from Regular Premium Balance": {
        "label": "Costi di investimento",
        "table": "T1",
        "pos": 3,
    },
    "Investment deduction from Single PremiumBalance": {
        "label": "Costi di investimento",
        "table": "T1",
        "pos": 3,
    },

    "Investment return from insurance funds": {
        "label": "Capitalizzazione",
        "table": "T1",
        "pos": 4,
    },

    "Investment return from insurance funds of Regular premium": {
        "label": "Capitalizzazione",
        "table": "T1",
        "pos": 4,
    },

    "Investment return from insurance funds of Single premium": {
        "label": "Capitalizzazione",
        "table": "T1",
        "pos": 4,
    },

    "Paid Premium": {
        "label": "Pagamenti dei Premi identificati",
        "table": "T1",
        "pos": 5,
    },
    
    "Paid Single Premium": {
        "label": "Pagamenti dei Premi identificati",
        "table": "T1",
        "pos": 5,
    },
  
    "Returned Premium": {
        "label": "Pagamenti dei Premi identificati",
        "table": "T1",
        "pos": 5,
    },

    "Risk deduction - Death": {
        "label": "Trattenuta copertura rischio morte",
        "table": "T1",
        "pos": 6,
    },
  
    "Risk deduction - accident insurance deduction": {
        "label": "Trattenuta copertura rischio infortunio",
        "table": "T1",
        "pos": 7,
    },
  
    "Risk deduction - Illnesses and operations": {
        "label": "Trattenuta copertura rischio malattia, interventi chirurgici e assistenza",
        "table": "T1",
        "pos": 8,
    },
  
   "Risk deduction - Waiver of premium": {
        "label": "Esonero Pagamento Premi ITP",
        "table": "T1",
        "pos": 9,      # appears after ordered rows
    },
  
    "Partial surrender": {
         "label": "Riscatto (parziale)",
         "table": "T1",
         "pos": 10,
    },
  
    "Partial Surrender Calculated value": {
         "label": "Riscatto (parziale)",
         "table": "T1",
         "pos": 10,
    },

    "Stamp Duty Fee": {
         "label": "Imposta di bollo",
         "table": "T1",
         "pos": 11,
    },

    # ────────────────  Table T2  ────────────────
    "Investment return of Novis Loyalty Bonus": {
        "label": "Rendimento Bonus Fedeltà NOVIS",
        "table": "T2",
        "pos": 1,
    },
    "Investment deduction of Novis Loyalty Bonus": {
        "label": "Costi di investimento - Bonus Fedeltà NOVIS",
        "table": "T2",
        "pos": 1,
    },
    # If the raw file already contains the Italian string use it directly:
    "NOVIS Loyalty Bonus": {
        "label": "Bonus Fedeltà NOVIS",
        "table": "T2",
        "pos": 2,
    },
  
    # ────────────────  Table T3 (Special Bonus)  ────────────────
    "NOVIS Special Bonus": {
        "label": "NOVIS Special Bonus",
        "table": "T3",
        "pos": 1,        # only row in its table
    },
   }
 
LABEL_POS = {cfg["label"]: cfg.get("pos", 999) for cfg in ITEM_CONFIG.values()}

POSITIVE_LABELS = {
    "Pagamenti dei Premi identificati",
    "Rendimento Bonus Fedeltà NOVIS",
    "Bonus Fedeltà NOVIS",
    "NOVIS Special Bonus",
}

TABLE_CONFIG = {
    # title empty → no "Item / Importo" header row (as in template)
    "T1": {"title": "", "include_total": True, "total_label": "Somma totale (escluso Bonus Fedeltà NOVIS e Special Bonus)"},
    "T2": {"title": "", "include_total": True, "total_label": "Bonus Fedeltà NOVIS con rendimento"},
    "T3": {"title": "", "include_total": False},
}

LETTER_SUBJECT_TPL = (
    "Dettaglio costi per il valore della Sua posizione assicurativa polizza n. "
    "{contract_number} al {calc_date_str} con codice fiscale {cf}."
)

OUTRO_PARAGRAPH = (
    "Qualora necessitasse di ulteriori informazioni in merito, La invitiamo "
    "gentilmente a riferirsi alla Tabella Costi contenuta nelle Condizioni di Assicurazione.\n\n"
    "Rimaniamo a disposizione per qualsiasi chiarimento e, ringraziando per la "
    "cortese attenzione, Le porgiamo i nostri più cordiali saluti."
)
GOODBYE_LINE = ""
SIGNATURE_BLOCK = (
    "Il team NOVIS"
)
# ── constants (add near the other CONFIG blocks) ──────────────────────
SALUTATION_ADDR = {
    "male": "Egr. Sig.",
    "female": "Gent.ma Sig.ra",
    "company": "Spett.le",
}
SALUTATION_GREET = {
    "male": "Egregio Signor",
    "female": "Gentilissima Signora",
    "company": "Spettabile",
}

COLUMN_ALIASES = {
    "EntryDate": "Item date",
    "ValueDate": "Item date",
    "EntryType": "Item name",
    "Amount": "Item value",
}
EXPECTED_COLS = {"Item date", "Item name", "Item value"}

# -------------------------------------------------------------------------
#  HELPERS
# -------------------------------------------------------------------------

def _fmt(amount: float) -> str:
    return format_currency(amount, "EUR", locale="it_IT")

def last_day_prev_month(d: date) -> date:
    """Return the last calendar day of the month preceding *d*."""
    return d.replace(day=1) - timedelta(days=1)

def last_name(name: str) -> str:
    """Return the surname; keep prefixes like 'Di', 'De', 'Del', etc."""
    tokens = name.split()
    if len(tokens) >= 2 and tokens[-2].lower() in {
        "di", "de", "del", "della", "d'", "da", "van", "von", "la", "le", "Di", "De", "Del", "Della", "D'", "Da", "Van", "Von", "La", "Le"
    }:
        return " ".join(tokens[-2:])           # 'Di Salvatore'
    return tokens[-1]                          # default: last token only

def make_intro(recipient_type: str, client_name: str, calc_date: str) -> str:
    """
    Builds the greeting + first paragraph in one shot, e.g.
    'Egregio Signor Rossi,\nsiamo con la presente … al 30/06/2025.'
    """
    if recipient_type == "company":
        greet_name = client_name
    else:
        greet_name = last_name(client_name)           # keeps 'Di Salvatore'

    greeting = f"{SALUTATION_GREET[recipient_type]} {greet_name},"
    body = (
        "siamo con la presente a trasmetterLe di seguito la tabella riportante il "
        "dettaglio dei costi applicati ai fini di calcolo del valore della Sua "
        f"posizione assicurativa al {calc_date}."
    )
    return f"{greeting}\n{body}"


def standardise_columns(df: pd.DataFrame) -> pd.DataFrame:
    if EXPECTED_COLS.issubset(df.columns):
        return df
    df = df.rename(columns={k: v for k, v in COLUMN_ALIASES.items() if k in df.columns})
    if not EXPECTED_COLS.issubset(df.columns):
        raise ValueError("Il file non contiene le colonne richieste.")
    return df


def aggregate_tables(df: pd.DataFrame) -> dict[str, pd.DataFrame]:
    df["Item value"] = pd.to_numeric(df["Item value"], errors="coerce")
    df = df.dropna(subset=["Item value"])
    df["Label"] = df["Item name"].apply(lambda x: ITEM_CONFIG.get(x, {}).get("label", x))
    df["Table"] = df["Item name"].apply(lambda x: ITEM_CONFIG.get(x, {}).get("table", "T1"))
    # invert sign for every label NOT in POSITIVE_LABELS (except we keep actual sign
    # for 'Capitalizzazione')
    df["Signed"] = df.apply(
      lambda r: r["Item value"]
      if r["Label"] in POSITIVE_LABELS or r["Label"] == "Capitalizzazione"
      else -r["Item value"],
      axis=1,
    )

    # create the grouped dataframe _before_ the loop
    grouped = (
        df.groupby(["Table", "Label"], as_index=False)["Signed"].sum()
        .rename(columns={"Signed": "Amount"})
    )
    grouped = grouped[grouped["Amount"] != 0]   # ← hide rows that net to €0

    tables = {}
    for tid, g in grouped.groupby("Table"):
      g["sort_key"] = g["Label"].apply(lambda x: LABEL_POS.get(x, 999))
      g = g.sort_values(["sort_key", "Label"]).drop(columns="sort_key")
      tables[tid] = g.drop(columns="Table")
    return tables

def doc_to_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def _safe_style(paragraph, style_name: str):
    """Apply Word style only if it exists in the template."""
    try:
        paragraph.style = style_name
    except KeyError:
        # template lacks the style – skip silently
        pass

def _safe_table_style(table, style_name: str) -> bool:
    try:
        table.style = style_name
        return True
    except KeyError:
        return False

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _add_thin_borders(tbl):
    """Ensure 0.5-pt black borders if the table style has none."""
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")     # 0.5 pt
        elem.set(qn("w:color"), "000000")

def parse_clipboard(blob: str) -> dict:
    """
    Extract contract, name, address, fiscal code from the pasted block.
    Returns {"contract": …, "name": …, "addr": …, "cf": …}
    Missing fields come back empty.
    """
    patterns = {
        "contract": r"Contract number:\s*(.+)",
        "name":     r"Policyholder:\s*(.+)",
        "addr":     r"Permanent residence:\s*(.+)",
        "cf":       r"Personal number:\s*(.+)",
    }
    out = {k: "" for k in patterns}
    for key, pat in patterns.items():
        m = re.search(pat, blob, flags=re.I)
        if m:
            out[key] = m.group(1).strip()
    return out

def split_addr(addr: str) -> list[str]:
    """
    Turn  'Street 8, 23849 Rogeno, Italy'
    into ['Street 8', '23849 Rogeno', 'Italy'].
    """
    return [p.strip() for p in addr.split(",") if p.strip()]


# -------------------------------------------------------------------------
#  DOCX BUILDER
# -------------------------------------------------------------------------

def build_doc(
    client_name: str,
    client_addr: str,
    cf: str,
    contract: str,
    calc_date: str,
    tables: dict[str, pd.DataFrame],
    recipient_type: str = "male",
    city: str = "Bratislava",
) -> Document:
    
    doc = Document("Novis_hl_papier_IT_motyl_12072023_prev.docx")
    doc.styles["Normal"].font.size = Pt(11)

    # address block
    prefix_short = SALUTATION_ADDR[recipient_type]
    p = doc.add_paragraph(f"{prefix_short} {client_name}")
    p.paragraph_format.left_indent = Inches(4)

    for line in split_addr(client_addr):
        q = doc.add_paragraph(line)
        q.paragraph_format.left_indent = Inches(4)

    today_str = date.today().strftime("%d/%m/%Y")
    r = doc.add_paragraph(f"{city}, {today_str}")
    doc.add_paragraph("")   # blank
    
    # 2-line replacement for the subject block
    p = doc.add_paragraph()
    _safe_style(p, "Heading 2")
    p.add_run("Dettaglio costi per il valore della Sua posizione assicurativa polizza n. ").bold = True
    p.add_run(contract).bold = True              # policy number
    p.add_run(f" al {calc_date} con codice fiscale ").bold = True
    p.add_run(cf).bold = True                    # codice fiscale
    
    doc.add_paragraph("")  # blank line after subject
    
    intro_text = make_intro(recipient_type, client_name, calc_date)
    para = doc.add_paragraph(intro_text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT      # greeting left-aligned
    doc.add_paragraph("")  # blank line after intro

    grand_total = 0
  
    # tables in predefined order
    for tid in [k for k in TABLE_CONFIG if k in tables]:
        cfg = TABLE_CONFIG[tid]
        df_tbl = tables[tid]

        if cfg["title"]:
            _safe_style(
                doc.add_paragraph(cfg["title"]),
                "Heading 3",
            )
      
        header = bool(cfg["title"])
        rows = 1 if header else 0
        tbl = doc.add_table(rows=rows, cols=2)   # create table
        if not _safe_table_style(tbl, "Table Grid"):
          _add_thin_borders(tbl)      # fallback when style is absent
          if header:
                  tbl.rows[0].cells[0].text = "Item"
                  hdr_imp = tbl.rows[0].cells[1]
                  hdr_imp.text = "Importo"
                  hdr_imp.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for _, row in df_tbl.iterrows():
            c1, c2 = tbl.add_row().cells
            c1.text = row["Label"]
            # bold the Special Bonus row
            if row["Label"] == "NOVIS Special Bonus":
               for run in c1.paragraphs[0].runs + c2.paragraphs[0].runs:
                    run.bold = True
            c2.text = _fmt(row["Amount"])
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        subtotal = df_tbl["Amount"].sum()
        if cfg.get("include_total"):
            c1, c2 = tbl.add_row().cells
            c1.text = cfg.get("total_label", "Totale")
            for r in c1.paragraphs[0].runs:
                r.bold = True
            c2.text = _fmt(subtotal)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            c2.paragraphs[0].runs[0].bold = True

        grand_total += subtotal
        doc.add_paragraph("")

    # grand-total as its own table
    gt = doc.add_table(rows=1, cols=2)
    _safe_table_style(gt, "Table Grid") or _add_thin_borders(gt)

    c1, c2 = gt.rows[0].cells
    c1.text = "Valore della Sua posizione assicurativa (incluso Bonus Fedeltà NOVIS e NOVIS Special Bonus)"
    c2.text = _fmt(grand_total)
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # bold label + amount
    for run in c1.paragraphs[0].runs:
      run.bold = True
    c2.paragraphs[0].runs[0].bold = True

    doc.add_paragraph("")           # spacer
    doc.add_paragraph(OUTRO_PARAGRAPH)
    doc.add_paragraph("")
    doc.add_paragraph(SIGNATURE_BLOCK)
    doc.add_paragraph("")                 # empty line after "Il team NOVIS"
    doc.add_paragraph(
    "NOVIS Insurance Company,\n"
    "NOVIS Versicherungsgesellschaft,\n"
    "NOVIS Compagnia di Assicurazioni,\n"
    "NOVIS Poisťovňa a.s."
    )
    return doc

# -------------------------------------------------------------------------
#  STREAMLIT FRONT END
# -------------------------------------------------------------------------

def main():
    st.set_page_config(page_title="📄 Generatore Lettera Valorizzazione",
                   layout="wide",  # more horizontal breathing room
                   initial_sidebar_state="expanded")    
    st.markdown(
    """
    <style>
    /* Center page & tighten width */
    .block-container { max-width: 1100px; padding-top: 1.5rem; }

    /* Sidebar title weight */
    [data-testid="stSidebar"] h1 { font-size:1.2rem; margin-bottom:0.5rem; }

    /* Table number alignment */
    td:nth-child(2) { text-align:right !important; font-variant-numeric: tabular-nums; }
    </style>
    """,
    unsafe_allow_html=True)

    file = st.file_uploader("Carica file movimenti (XLS/XLSX)", type=["xls", "xlsx"])

    st.subheader("Dati cliente")
    
    # --- new clipboard import UI -------------------------------------------
    st.subheader("Incolla dati dal sistema interno")
    clip_txt = st.text_area("Blocca-dati", height=140, key="clip")

    if st.button("→ Importa") and clip_txt.strip():
        parsed = parse_clipboard(clip_txt)
        parsed["addr"] = "\n".join(split_addr(parsed["addr"]))   # ↵ inside widget
        st.session_state.update(
            contract=parsed["contract"],
            name=parsed["name"],
            addr=parsed["addr"],
            cf=parsed["cf"],
        )

    # --- recipient selector --------------------------------------------------
    label2value = {"Uomo": "male", "Donna": "female", "Società": "company"}
    recip_label = st.selectbox("Destinatario", list(label2value.keys()))
    recipient_type = label2value[recip_label]

    # --- city input (so the variable exists) ---------------------------------
    city = st.text_input("Luogo (prefisso alla data)", "Bratislava")

  
    name = st.text_input("Nome", key="name")
    addr = st.text_area("Indirizzo", key="addr")
    cf = st.text_input("Codice fiscale", key="cf")
    contract = st.text_input("Numero polizza", key="contract")
    # --- valore al  fine‑mese -------------------------------------------------
    # build a list of month‑end dates: previous month first, then going backwards
    today = date.today()
    opts = []
    d = last_day_prev_month(today)       # default = last day of previous month
    for _ in range(12):                  # make e.g. 3 years of choices
        opts.append(d)
        d = last_day_prev_month(d)       # hop back one more month

    calc_date = st.selectbox(
        "Data valorizzazione (solo fine mese)",
        opts,
        index=0,                         # pre‑select previous month end
        format_func=lambda x: x.strftime("%d/%m/%Y"),
    )
    calc_date_str = calc_date.strftime("%d/%m/%Y")   # string for the DOCX

    if file is not None:
        try:
            df = standardise_columns(pd.read_excel(file))
        except Exception as e:
            st.error(f"Errore nel file: {e}")
            st.stop()
        tables = aggregate_tables(df)
        st.subheader("Anteprima")
        for tid in [k for k in TABLE_CONFIG if k in tables]:
            tbl_df = tables[tid]
            cfg = TABLE_CONFIG[tid]
            st.markdown(f"#### {cfg['title'] or 'Tabella costi'}")
            st.dataframe(tbl_df.assign(Importo=tbl_df["Amount"].apply(_fmt)).drop(columns="Amount"), use_container_width=True)
            if cfg.get("include_total"):
                st.markdown(f"**{cfg['total_label']}: {_fmt(tbl_df['Amount'].sum())}**")

        if all([name, addr, cf, contract]):
            doc = build_doc(
              name, addr, cf, contract, calc_date_str, tables,
              recipient_type=recipient_type,
              city=city,
            )
            st.download_button(
                label="⬇️ Scarica Word",
                data=doc_to_bytes(doc),
                file_name=f"VAL_{contract}_{date.today().strftime('%d%m%y')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             ) 
        else:
            st.info("Compila tutti i campi cliente per generare la lettera.")


if __name__ == "__main__":
    main()

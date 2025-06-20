# streamlit_insurance_letter.py
"""
Streamlit app to generate a summary letter for insurance contract movements.
Upload an XLS/XLSX file containing columns:

- Item date
- Item name
- Item value

The app groups rows by **Item name**, sums **Item value**, translates each
item name into another language (using a hard‑coded dictionary), and
injects the results into a Word document letter ready to send to the client.
"""

# ---- Imports --------------------------------------------------------------
import streamlit as st
import pandas as pd
from datetime import date
from io import BytesIO
from docx import Document
import streamlit.components.v1 as components


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  HARD‑CODED CONFIGURATION                                            ║
# ║  ▸ Edit the values in this block to localise the app                 ║
# ║    (translations, letter subject/body, sign‑off, etc.).              ║
# ╚══════════════════════════════════════════════════════════════════════╝

TRANSLATION_MAP = {
    # "Original Item Name": "Translated Item Name",
    "Premium": "Poistné",
    "Claim": "Škoda",
    "Commission": "Provízia",
    # ▸ Add more translations as needed …
}

LETTER_SUBJECT = "Statement of Account – Insurance Contract"

LETTER_BODY_HEADER = (
    "Dear {client_name},\n\n"
    "Please find below the statement of movements on your insurance contract "
    "number {contract_number}."
)

GOODBYE_LINE = "Sincerely,"

SIGNATURE_BLOCK = "Your Insurance Company\nInsurance Operations Team"

# ──────────────────────────────────────────────────────────────────────────
#  END OF HARD‑CODED SECTION                                               
# ──────────────────────────────────────────────────────────────────────────


def summarise_data(df: pd.DataFrame, translation: dict) -> pd.DataFrame:
    """Group by *Item name*, sum *Item value*, and translate the item name."""
    summary = df.groupby("Item name", as_index=False)["Item value"].sum()
    summary["Item"] = summary["Item name"].map(translation).fillna(summary["Item name"])
    summary.rename(columns={"Item value": "Amount"}, inplace=True)
    return summary[["Item", "Amount"]]


def build_letter_text(
    client_name: str,
    client_address: str,
    letter_date: date,
    contract_number: str,
    summary: pd.DataFrame,
) -> str:
    """Return the plain‑text version of the letter (used for copy‑&‑paste)."""
    lines: list[str] = []
    lines.append(client_name)
    lines.append(client_address)
    lines.append("")
    lines.append(letter_date.strftime("%d %B %Y"))
    lines.append("")
    lines.append(LETTER_SUBJECT)
    lines.append("")
    lines.append(LETTER_BODY_HEADER.format(
        client_name=client_name, contract_number=contract_number
    ))
    lines.append("")
    lines.append(summary.to_string(index=False, header=True))
    lines.append("")
    lines.append(GOODBYE_LINE)
    lines.append("")
    lines.append(SIGNATURE_BLOCK)
    return "\n".join(lines)


def build_letter_doc(
    client_name: str,
    client_address: str,
    letter_date: date,
    contract_number: str,
    summary: pd.DataFrame,
) -> Document:
    """Return a python‑docx Document ready to be saved/downloaded."""
    doc = Document()

    # Header block
    doc.add_paragraph(client_name)
    doc.add_paragraph(client_address)
    doc.add_paragraph("")
    doc.add_paragraph(letter_date.strftime("%d %B %Y"))
    doc.add_paragraph("")

    # Subject line
    doc.add_paragraph(LETTER_SUBJECT).style = "Heading 2"
    doc.add_paragraph("")

    # Body
    doc.add_paragraph(
        LETTER_BODY_HEADER.format(client_name=client_name, contract_number=contract_number)
    )

    # Table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Amount"

    for _, row in summary.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Item"])
        row_cells[1].text = f"{row['Amount']:.2f}"

    doc.add_paragraph("")
    doc.add_paragraph(GOODBYE_LINE)
    doc.add_paragraph("")
    doc.add_paragraph(SIGNATURE_BLOCK)
    return doc


def doc_to_bytes(doc: Document) -> BytesIO:
    """Serialize a python‑docx Document to memory."""
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main() -> None:
    st.set_page_config(
        page_title="Insurance Letter Generator",
        layout="centered",
    )
    st.title("📄 Insurance Contract Letter Generator")

    # 1️⃣ Upload Excel
    uploaded_file = st.file_uploader(
        "Upload XLS/XLSX file with movements", type=["xls", "xlsx"]
    )

    # 2️⃣ Client info
    st.subheader("Client Information")
    client_name = st.text_input("Client Name")
    contract_number = st.text_input("Contract Number")
    client_address = st.text_area("Client Address")
    letter_date = st.date_input("Letter Date", value=date.today())

    if uploaded_file is not None:
        # Validate & preview data
        try:
            df = pd.read_excel(uploaded_file)
        except Exception as err:
            st.error(f"❌ Could not read the file: {err}")
            st.stop()

        expected_cols = {"Item date", "Item name", "Item value"}
        if not expected_cols.issubset(df.columns):
            st.error(
                "❌ The uploaded file must contain the columns: "
                + ", ".join(expected_cols)
            )
            st.stop()

        summary_df = summarise_data(df, TRANSLATION_MAP)

        st.subheader("Summarised Movements (translated)")
        st.dataframe(summary_df, use_container_width=True)

        # Generate letter once mandatory fields are filled
        if all([client_name, contract_number, client_address]):
            letter_txt = build_letter_text(
                client_name,
                client_address,
                letter_date,
                contract_number,
                summary_df,
            )

            # Preview & copy
            st.subheader("Letter Preview & Copy")
            html_block = f"""
            <textarea id=\"letterArea\" style=\"width:100%;height:260px;\">{letter_txt}</textarea><br>
            <button style=\"margin-top:6px;padding:6px 12px;font-size:14px;\" onclick=\"navigator.clipboard.writeText(document.getElementById('letterArea').value)\">📋 Copy Letter</button>
            """
            components.html(html_block, height=320)

            # Build Word doc & offer download
            doc = build_letter_doc(
                client_name,
                client_address,
                letter_date,
                contract_number,
                summary_df,
            )
            st.download_button(
                label="⬇️ Download Word Document",
                data=doc_to_bytes(doc),
                file_name="Insurance_Statement.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.info("ℹ️ Please fill in *all* client information fields to generate the letter.")


if __name__ == "__main__":
    main()
